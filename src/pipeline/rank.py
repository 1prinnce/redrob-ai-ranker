"""Online candidate ranking pipeline.

This script loads offline preprocessing artifacts, parses a job description,
retrieves semantically similar candidates with FAISS, scores only the retrieved
pool with cached static features, and writes the final top-100 submission CSV.
"""

from __future__ import annotations

from collections.abc import Iterator, Mapping, Sequence
from contextlib import contextmanager
from dataclasses import dataclass
from numbers import Integral, Real
import argparse
import csv
import json
import os
import pathlib
import pickle
import sys
import time
from typing import Any, Final, TypedDict

import numpy as np
import numpy.typing as npt


os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")
os.environ.setdefault("HF_DATASETS_OFFLINE", "1")

if __package__ in {None, ""}:  # pragma: no cover - supports direct script execution.
    sys.path.append(str(pathlib.Path(__file__).resolve().parents[2]))

from src.embeddings import encoder
from src.embeddings import index as faiss_index
from src.jd.parser import ParsedJD, extract_requirements_text, parse_jd
from src.scoring import normalizer
from src.scoring.reasoning import generate_reasoning, validate_reasoning_variance
from src.scoring.scorer import (
    BEHAVIORAL_WEIGHT,
    CAREER_WEIGHT,
    FALLBACK_SCORE,
    JD_MATCH_WEIGHT,
    LOCATION_WEIGHT,
    SKILLS_WEIGHT,
    safe_score,
)
from src.utils.logger import get_logger, tqdm_loader
from src.utils.timer import BudgetTimer


LOGGER = get_logger(__name__)

DEFAULT_MODEL_NAME: Final[str] = "all-MiniLM-L6-v2"
MAX_RUNTIME_SECONDS: Final[float] = 300.0
RETRIEVAL_K: Final[int] = 1000
TOP_N: Final[int] = 100
RETRIEVAL_TARGET_SECONDS: Final[float] = 1.0
SCORE_DECIMAL_PLACES: Final[int] = 6

EMBEDDINGS_FILENAME: Final[str] = "candidate_embeddings.npy"
FEATURES_FILENAME: Final[str] = "candidate_features.npz"
METADATA_FILENAME: Final[str] = "candidate_meta.pkl"
FAISS_INDEX_FILENAME: Final[str] = "faiss.index"
MANIFEST_FILENAME: Final[str] = "prep_manifest.json"

FEATURE_NAMES: Final[tuple[str, ...]] = (
    "career_score",
    "behavioral_score",
    "location_score",
    "skill_score",
    "honeypot_penalty",
    "disqualifier_penalty",
)
METADATA_FIELDS: Final[tuple[str, ...]] = (
    "candidate_id",
    "current_title",
    "current_company",
    "years_of_experience",
    "notice_period_days",
    "recruiter_response_rate",
)
CSV_COLUMNS: Final[tuple[str, ...]] = ("candidate_id", "rank", "score", "reasoning")


class CandidateMetaCache(TypedDict):
    """Columnar candidate metadata cache aligned to embedding row order."""

    candidate_id: list[str]
    current_title: list[str | None]
    current_company: list[str | None]
    years_of_experience: list[float | None]
    notice_period_days: list[float | None]
    recruiter_response_rate: list[float | None]


class SubmissionRow(TypedDict):
    """One row in the final submission CSV."""

    candidate_id: str
    rank: int
    score: float
    reasoning: str


@dataclass(frozen=True)
class RankArtifactPaths:
    """Resolved artifact paths required by the online ranker."""

    embeddings: pathlib.Path
    features: pathlib.Path
    metadata: pathlib.Path
    faiss_index: pathlib.Path
    manifest: pathlib.Path


@dataclass(frozen=True)
class LoadedArtifacts:
    """Loaded and validated ranking artifacts."""

    paths: RankArtifactPaths
    index: Any
    embeddings: np.ndarray
    features: dict[str, npt.NDArray[np.float32]]
    metadata: CandidateMetaCache
    manifest: Mapping[str, Any]
    candidate_count: int
    embedding_dimension: int
    model_name: str


@dataclass(frozen=True)
class JDContext:
    """Parsed JD intent and dense text used for embedding retrieval."""

    parsed: ParsedJD
    requirements_text: str


@dataclass(frozen=True)
class ScoredCandidate:
    """A retrieved candidate with full online scoring diagnostics."""

    candidate_index: int
    candidate_id: str
    raw_score: float
    semantic_similarity: float
    breakdown: dict[str, float]
    candidate_record: dict[str, Any]


def main(argv: Sequence[str] | None = None) -> int:
    """Run the online ranking pipeline from command-line arguments.

    Args:
        argv: Optional argument vector used by tests. ``None`` reads from
            ``sys.argv`` through ``argparse``.

    Returns:
        Process-style exit code: ``0`` on success, ``1`` on failure.
    """
    args = _parse_args(argv)
    artifacts_dir = args.artifacts.expanduser().resolve()
    jd_path = args.jd.expanduser().resolve()
    out_path = args.out.expanduser().resolve()
    timer = BudgetTimer(max_seconds=MAX_RUNTIME_SECONDS)

    try:
        _validate_cli_inputs(artifacts_dir, jd_path)

        with _stage_timer("load_artifacts", timer):
            artifacts = _load_artifacts(artifacts_dir)

        with _stage_timer("parse_jd", timer):
            jd_context = _parse_job_description(jd_path)

        with _stage_timer("encode_jd", timer):
            jd_embedding = _encode_jd(
                jd_context.requirements_text,
                model_name=artifacts.model_name,
                expected_dimension=artifacts.embedding_dimension,
            )

        with _stage_timer("semantic_retrieval", timer):
            retrieved_indices, similarities = _retrieve_candidates(
                artifacts.index,
                jd_embedding,
                candidate_count=artifacts.candidate_count,
            )

        with _stage_timer("candidate_scoring", timer):
            scored_candidates = _score_candidates(
                artifacts=artifacts,
                jd_embedding=jd_embedding,
                retrieved_indices=retrieved_indices,
                similarities=similarities,
            )

        with _stage_timer("ranking", timer):
            top_candidates, normalized_scores = _rank_top_candidates(scored_candidates)

        with _stage_timer("reasoning_generation", timer):
            submission_rows = _build_submission_rows(top_candidates, normalized_scores)

        with _stage_timer("submission_validation", timer):
            _validate_submission(submission_rows)

        with _stage_timer("write_csv", timer):
            _write_submission(submission_rows, out_path)

        LOGGER.info(
            "Online ranking complete",
            extra={
                "candidate_count": artifacts.candidate_count,
                "retrieved_count": len(retrieved_indices),
                "submitted_count": len(submission_rows),
                "out_path": str(out_path),
                "elapsed_seconds": round(timer.elapsed(), 3),
            },
        )
        return 0
    except Exception as exc:
        LOGGER.exception("Online ranking failed: %s", exc)
        return 1


def _parse_args(argv: Sequence[str] | None) -> argparse.Namespace:
    """Parse online ranking CLI arguments."""
    parser = argparse.ArgumentParser(description="Rank candidates for a job description.")
    parser.add_argument(
        "--artifacts",
        required=True,
        type=pathlib.Path,
        help="Directory containing preprocessing artifacts.",
    )
    parser.add_argument(
        "--jd",
        required=True,
        type=pathlib.Path,
        help="Path to the job description text file.",
    )
    parser.add_argument(
        "--out",
        required=True,
        type=pathlib.Path,
        help="Path where submission.csv will be written.",
    )
    return parser.parse_args(argv)


def _validate_cli_inputs(artifacts_dir: pathlib.Path, jd_path: pathlib.Path) -> None:
    """Validate cheap CLI inputs before expensive model and FAISS work."""
    if not artifacts_dir.exists():
        raise FileNotFoundError(f"Artifact directory does not exist: {artifacts_dir}")
    if not artifacts_dir.is_dir():
        raise ValueError(f"Artifact path is not a directory: {artifacts_dir}")
    if not jd_path.exists():
        raise FileNotFoundError(f"Job description file does not exist: {jd_path}")
    if not jd_path.is_file():
        raise ValueError(f"Job description path is not a file: {jd_path}")


@contextmanager
def _stage_timer(stage_name: str, timer: BudgetTimer) -> Iterator[None]:
    """Log one stage duration and check the remaining runtime budget."""
    LOGGER.info("Starting stage: %s", stage_name)
    start = time.perf_counter()
    try:
        yield
    except Exception:
        elapsed = time.perf_counter() - start
        LOGGER.exception("Stage failed: %s", stage_name, extra={"duration_seconds": round(elapsed, 3)})
        raise
    else:
        elapsed = time.perf_counter() - start
        LOGGER.info("Finished stage: %s", stage_name, extra={"duration_seconds": round(elapsed, 3)})
        timer.check(f"after {stage_name}")


def _artifact_paths(artifacts_dir: pathlib.Path) -> RankArtifactPaths:
    """Return expected artifact paths for an artifact directory."""
    return RankArtifactPaths(
        embeddings=artifacts_dir / EMBEDDINGS_FILENAME,
        features=artifacts_dir / FEATURES_FILENAME,
        metadata=artifacts_dir / METADATA_FILENAME,
        faiss_index=artifacts_dir / FAISS_INDEX_FILENAME,
        manifest=artifacts_dir / MANIFEST_FILENAME,
    )


def _load_artifacts(artifacts_dir: pathlib.Path) -> LoadedArtifacts:
    """Load all ranking artifacts and validate cross-file alignment."""
    paths = _artifact_paths(artifacts_dir)
    _validate_artifact_files(paths)

    manifest = _load_manifest(paths.manifest)
    retrieval_index = faiss_index.load_index(paths.faiss_index)
    embeddings = _load_embedding_memmap(paths.embeddings)
    features = _load_feature_arrays(paths.features)
    metadata = _load_metadata_cache(paths.metadata)
    candidate_count, embedding_dimension = _validate_artifact_alignment(
        index=retrieval_index,
        embeddings=embeddings,
        features=features,
        metadata=metadata,
        manifest=manifest,
    )
    model_name = _manifest_model_name(manifest)

    LOGGER.info(
        "Loaded ranking artifacts",
        extra={
            "candidate_count": candidate_count,
            "embedding_dimension": embedding_dimension,
            "model_name": model_name,
        },
    )
    return LoadedArtifacts(
        paths=paths,
        index=retrieval_index,
        embeddings=embeddings,
        features=features,
        metadata=metadata,
        manifest=manifest,
        candidate_count=candidate_count,
        embedding_dimension=embedding_dimension,
        model_name=model_name,
    )


def _validate_artifact_files(paths: RankArtifactPaths) -> None:
    """Validate that every required artifact file exists and is non-empty."""
    for artifact_name, path in (
        ("candidate_embeddings", paths.embeddings),
        ("candidate_features", paths.features),
        ("candidate_meta", paths.metadata),
        ("faiss_index", paths.faiss_index),
        ("prep_manifest", paths.manifest),
    ):
        if not path.exists():
            raise FileNotFoundError(f"Missing {artifact_name} artifact: {path}")
        if not path.is_file():
            raise ValueError(f"{artifact_name} artifact is not a file: {path}")
        if path.stat().st_size <= 0:
            raise ValueError(f"{artifact_name} artifact is empty: {path}")


def _load_manifest(path: pathlib.Path) -> Mapping[str, Any]:
    """Load and validate the preprocessing manifest."""
    try:
        with path.open("r", encoding="utf-8") as handle:
            value = json.load(handle)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Prep manifest is not valid JSON: {path}: {exc}") from exc

    if not isinstance(value, Mapping):
        raise TypeError(f"Prep manifest must contain a JSON object, got {type(value).__name__}.")
    return value


def _load_embedding_memmap(path: pathlib.Path) -> np.ndarray:
    """Load candidate embeddings as a read-only memory map."""
    if path.suffix != ".npy":
        raise ValueError(f"Candidate embeddings must be stored in a .npy file: {path}")
    try:
        embeddings = np.load(path, mmap_mode="r", allow_pickle=False)
    except OSError as exc:
        raise OSError(f"Unable to load candidate embeddings from {path}: {exc}") from exc

    if not isinstance(embeddings, np.ndarray):
        raise TypeError(f"Candidate embeddings must load as a numpy array, got {type(embeddings).__name__}.")
    if embeddings.dtype != np.float32:
        raise TypeError(f"Candidate embeddings must be float32, got {embeddings.dtype}.")
    if embeddings.ndim != 2:
        raise ValueError(f"Candidate embeddings must be two-dimensional, got shape {embeddings.shape}.")
    if embeddings.shape[0] <= 0 or embeddings.shape[1] <= 0:
        raise ValueError(f"Candidate embeddings must be non-empty, got shape {embeddings.shape}.")
    return embeddings


def _load_feature_arrays(path: pathlib.Path) -> dict[str, npt.NDArray[np.float32]]:
    """Load cached static candidate feature arrays."""
    try:
        archive = np.load(path, allow_pickle=False)
    except OSError as exc:
        raise OSError(f"Unable to load candidate features from {path}: {exc}") from exc

    features: dict[str, npt.NDArray[np.float32]] = {}
    with archive:
        missing = [name for name in FEATURE_NAMES if name not in archive.files]
        if missing:
            raise ValueError(f"Candidate features missing arrays: {', '.join(missing)}")

        for name in FEATURE_NAMES:
            array = np.asarray(archive[name], dtype=np.float32)
            if array.ndim != 1:
                raise ValueError(f"Feature {name!r} must be one-dimensional, got shape {array.shape}.")
            if array.size == 0:
                raise ValueError(f"Feature {name!r} must not be empty.")
            if not np.isfinite(array).all():
                raise ValueError(f"Feature {name!r} contains NaN or infinite values.")
            features[name] = array

    return features


def _load_metadata_cache(path: pathlib.Path) -> CandidateMetaCache:
    """Load the candidate metadata cache and validate column alignment."""
    try:
        with path.open("rb") as handle:
            value = pickle.load(handle)
    except (EOFError, OSError, pickle.UnpicklingError) as exc:
        raise ValueError(f"Unable to load candidate metadata cache from {path}: {exc}") from exc

    if not isinstance(value, Mapping):
        raise TypeError(f"Candidate metadata cache must be a mapping, got {type(value).__name__}.")

    cache: CandidateMetaCache = {
        "candidate_id": _as_list_field(value, "candidate_id"),
        "current_title": _as_list_field(value, "current_title"),
        "current_company": _as_list_field(value, "current_company"),
        "years_of_experience": _as_list_field(value, "years_of_experience"),
        "notice_period_days": _as_list_field(value, "notice_period_days"),
        "recruiter_response_rate": _as_list_field(value, "recruiter_response_rate"),
    }
    _validate_metadata_cache(cache)
    return cache


def _as_list_field(value: Mapping[str, Any], field_name: str) -> list[Any]:
    """Return one metadata field as a list while rejecting scalar strings."""
    if field_name not in value:
        raise ValueError(f"Candidate metadata cache missing field: {field_name}")

    field_value = value[field_name]
    if isinstance(field_value, (str, bytes)) or not isinstance(field_value, Sequence):
        raise TypeError(f"Metadata field {field_name!r} must be a sequence.")
    return list(field_value)


def _validate_metadata_cache(cache: CandidateMetaCache) -> None:
    """Validate metadata cache lengths, candidate IDs, and numeric fields."""
    lengths = {field: len(cache[field]) for field in METADATA_FIELDS}
    if len(set(lengths.values())) != 1:
        raise ValueError(f"Metadata fields have inconsistent lengths: {lengths}")
    if lengths["candidate_id"] == 0:
        raise ValueError("Candidate metadata cache must not be empty.")

    seen_ids: set[str] = set()
    for index, candidate_id in enumerate(cache["candidate_id"]):
        if not isinstance(candidate_id, str) or not candidate_id:
            raise ValueError(f"candidate_id[{index}] must be a non-empty string.")
        if candidate_id in seen_ids:
            raise ValueError(f"Duplicate candidate_id in metadata cache: {candidate_id!r}")
        seen_ids.add(candidate_id)

    for field_name in ("years_of_experience", "notice_period_days", "recruiter_response_rate"):
        for index, value in enumerate(cache[field_name]):
            if value is not None and not _is_finite_real(value):
                raise ValueError(f"{field_name}[{index}] must be a finite number or None.")


def _validate_artifact_alignment(
    *,
    index: Any,
    embeddings: np.ndarray,
    features: Mapping[str, npt.NDArray[np.float32]],
    metadata: CandidateMetaCache,
    manifest: Mapping[str, Any],
) -> tuple[int, int]:
    """Validate candidate counts and embedding dimensions across artifacts."""
    candidate_count = len(metadata["candidate_id"])
    embedding_dimension = int(embeddings.shape[1])

    if candidate_count < TOP_N:
        raise ValueError(f"At least {TOP_N} candidates are required, got {candidate_count}.")
    if int(index.ntotal) != candidate_count:
        raise ValueError(f"FAISS index count {int(index.ntotal)} does not match metadata count {candidate_count}.")
    if int(index.d) != embedding_dimension:
        raise ValueError(f"FAISS dimension {int(index.d)} does not match embeddings dimension {embedding_dimension}.")
    if int(embeddings.shape[0]) != candidate_count:
        raise ValueError(
            f"Embedding row count {int(embeddings.shape[0])} does not match metadata count {candidate_count}."
        )

    for name in FEATURE_NAMES:
        feature_count = int(features[name].shape[0])
        if feature_count != candidate_count:
            raise ValueError(f"Feature {name!r} count {feature_count} does not match {candidate_count}.")

    manifest_count = _required_manifest_int(manifest, "candidate_count")
    manifest_dimension = _required_manifest_int(manifest, "embedding_dimension")
    if manifest_count != candidate_count:
        raise ValueError(f"Manifest candidate_count {manifest_count} does not match {candidate_count}.")
    if manifest_dimension != embedding_dimension:
        raise ValueError(f"Manifest embedding_dimension {manifest_dimension} does not match {embedding_dimension}.")

    manifest_features = manifest.get("feature_names")
    if isinstance(manifest_features, Sequence) and not isinstance(manifest_features, (str, bytes)):
        manifest_feature_names = {item for item in manifest_features if isinstance(item, str)}
        missing_features = [name for name in FEATURE_NAMES if name not in manifest_feature_names]
        if missing_features:
            raise ValueError(f"Manifest feature_names missing: {', '.join(missing_features)}")

    return candidate_count, embedding_dimension


def _required_manifest_int(manifest: Mapping[str, Any], field_name: str) -> int:
    """Return a positive integer field from the manifest."""
    value = manifest.get(field_name)
    if isinstance(value, bool) or not isinstance(value, Integral):
        raise TypeError(f"Manifest field {field_name!r} must be an integer, got {type(value).__name__}.")
    clean_value = int(value)
    if clean_value <= 0:
        raise ValueError(f"Manifest field {field_name!r} must be positive, got {clean_value}.")
    return clean_value


def _manifest_model_name(manifest: Mapping[str, Any]) -> str:
    """Return the embedding model name recorded by preprocessing."""
    value = manifest.get("model_name")
    if isinstance(value, str) and value.strip():
        return value.strip()
    LOGGER.warning("Prep manifest missing model_name; falling back to %s", DEFAULT_MODEL_NAME)
    return DEFAULT_MODEL_NAME


def _parse_job_description(jd_path: pathlib.Path) -> JDContext:
    """Read a TXT or DOCX job description into structured and dense forms."""
    suffix = jd_path.suffix.casefold()
    LOGGER.info("Reading job description", extra={"path": str(jd_path), "file_type": suffix or "<none>"})

    if suffix == ".txt":
        try:
            jd_text = jd_path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError(f"Job description must be valid UTF-8 text: {jd_path}: {exc}") from exc
    elif suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to read .docx job descriptions.") from exc

        try:
            document = Document(str(jd_path))
        except Exception as exc:
            raise ValueError(f"Unable to read DOCX job description: {jd_path}: {exc}") from exc

        jd_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        LOGGER.info(
            "Extracted DOCX job description text",
            extra={"path": str(jd_path), "paragraph_count": len(document.paragraphs), "char_count": len(jd_text)},
        )
    else:
        raise ValueError(f"Unsupported job description file type {suffix!r}; expected .txt or .docx.")

    parsed = parse_jd(jd_text)
    requirements_text = extract_requirements_text(jd_text).strip()
    if not requirements_text:
        requirements_text = " ".join(jd_text.split())
    if not requirements_text:
        raise ValueError(f"Job description has no usable ranking text: {jd_path}")

    LOGGER.info(
        "Parsed job description",
        extra={
            "must_haves": parsed["must_haves"],
            "nice_to_haves": parsed["nice_to_haves"],
            "experience_range": parsed["experience_range"],
            "preferred_locations": parsed["preferred_locations"],
            "domain_focus": parsed["domain_focus"],
            "requirements_chars": len(requirements_text),
        },
    )
    return JDContext(parsed=parsed, requirements_text=requirements_text)


def _encode_jd(requirements_text: str, *, model_name: str, expected_dimension: int) -> npt.NDArray[np.float32]:
    """Load the embedding model and generate one normalized JD embedding."""
    model = encoder.load_model(model_name)
    embedding = encoder.encode_single(requirements_text, model=model)
    if embedding.dtype != np.float32:
        embedding = np.asarray(embedding, dtype=np.float32)
    if embedding.ndim != 1:
        raise ValueError(f"JD embedding must be one-dimensional, got shape {embedding.shape}.")
    if int(embedding.shape[0]) != expected_dimension:
        raise ValueError(
            f"JD embedding dimension {int(embedding.shape[0])} does not match artifacts dimension {expected_dimension}."
        )
    if not np.isfinite(embedding).all():
        raise ValueError("JD embedding contains NaN or infinite values.")

    norm = float(np.linalg.norm(embedding))
    if not np.isclose(norm, 1.0, atol=1e-4):
        raise ValueError(f"JD embedding is not L2-normalized: norm={norm:.6f}.")
    return np.ascontiguousarray(embedding, dtype=np.float32)


def _retrieve_candidates(
    index: Any,
    jd_embedding: npt.NDArray[np.float32],
    *,
    candidate_count: int,
) -> tuple[npt.NDArray[np.int64], npt.NDArray[np.float32]]:
    """Retrieve top semantic candidates from FAISS."""
    retrieval_start = time.perf_counter()
    distances, indices = faiss_index.query_index(index, jd_embedding, k=RETRIEVAL_K)
    retrieval_seconds = time.perf_counter() - retrieval_start

    flat_indices = np.asarray(indices, dtype=np.int64).reshape(-1)
    flat_similarities = np.asarray(distances, dtype=np.float32).reshape(-1)
    valid_mask = flat_indices >= 0
    flat_indices = flat_indices[valid_mask]
    flat_similarities = flat_similarities[valid_mask]

    if flat_indices.size < TOP_N:
        raise ValueError(f"FAISS returned only {flat_indices.size} valid candidates; {TOP_N} are required.")
    if np.any(flat_indices >= candidate_count):
        bad_index = int(flat_indices[np.argmax(flat_indices >= candidate_count)])
        raise ValueError(f"FAISS returned out-of-range candidate index {bad_index} for count {candidate_count}.")
    if len(set(int(index_value) for index_value in flat_indices)) != int(flat_indices.size):
        raise ValueError("FAISS retrieval returned duplicate candidate indices.")
    if not np.isfinite(flat_similarities).all():
        raise ValueError("FAISS retrieval returned NaN or infinite similarities.")

    log_payload = {
        "retrieved_count": int(flat_indices.size),
        "latency_seconds": round(retrieval_seconds, 4),
        "target_seconds": RETRIEVAL_TARGET_SECONDS,
    }
    if retrieval_seconds > RETRIEVAL_TARGET_SECONDS:
        LOGGER.warning("Semantic retrieval exceeded latency target", extra=log_payload)
    else:
        LOGGER.info("Semantic retrieval met latency target", extra=log_payload)

    return flat_indices, flat_similarities


def _score_candidates(
    *,
    artifacts: LoadedArtifacts,
    jd_embedding: npt.NDArray[np.float32],
    retrieved_indices: npt.NDArray[np.int64],
    similarities: npt.NDArray[np.float32],
) -> list[ScoredCandidate]:
    """Score the retrieved candidates using cached features and safe fallbacks."""
    if retrieved_indices.shape != similarities.shape:
        raise ValueError(
            "retrieved_indices and similarities must have matching shapes; "
            f"got {retrieved_indices.shape} and {similarities.shape}."
        )

    scored: list[ScoredCandidate] = []
    pairs = list(zip(retrieved_indices.tolist(), similarities.tolist(), strict=True))
    for candidate_index, similarity in tqdm_loader(pairs, total=len(pairs), desc="Scoring candidates"):
        row_index = int(candidate_index)
        candidate_record = _candidate_record_from_metadata(artifacts.metadata, row_index)
        candidate_embedding = np.asarray(artifacts.embeddings[row_index], dtype=np.float32)
        guarded_score = safe_score(candidate_record, jd_embedding, candidate_embedding, idx=row_index)

        try:
            breakdown = _cached_score_breakdown(
                features=artifacts.features,
                candidate_index=row_index,
                semantic_similarity=float(similarity),
                guarded_score=guarded_score,
            )
        except Exception as exc:
            LOGGER.exception(
                "Cached scoring failed; using safe_score fallback",
                extra={"candidate_index": row_index, "candidate_id": candidate_record["candidate_id"], "error": str(exc)},
            )
            fallback = _clean_score(guarded_score, default=FALLBACK_SCORE)
            breakdown = {
                "semantic_similarity": _clean_score(similarity, default=0.0),
                "career_score": 0.0,
                "behavioral_score": 0.0,
                "location_score": 0.0,
                "skill_score": 0.0,
                "honeypot_penalty": 0.0,
                "disqualifier_penalty": 0.0,
                "base_score": fallback,
                "final_score": fallback,
                "safe_score": fallback,
            }

        scored.append(
            ScoredCandidate(
                candidate_index=row_index,
                candidate_id=str(candidate_record["candidate_id"]),
                raw_score=breakdown["final_score"],
                semantic_similarity=breakdown["semantic_similarity"],
                breakdown=breakdown,
                candidate_record=candidate_record,
            )
        )

    _validate_scored_candidates(scored)
    LOGGER.info("Scored retrieved candidates", extra={"candidate_count": len(scored)})
    return scored


def _cached_score_breakdown(
    *,
    features: Mapping[str, npt.NDArray[np.float32]],
    candidate_index: int,
    semantic_similarity: float,
    guarded_score: float,
) -> dict[str, float]:
    """Combine semantic similarity with cached static features and penalties."""
    semantic_score = _clamp(_finite_float(semantic_similarity, "semantic_similarity"))
    career_score = _clamp(_feature_value(features, "career_score", candidate_index))
    behavioral_score = _clamp(_feature_value(features, "behavioral_score", candidate_index))
    location_score = _clamp(_feature_value(features, "location_score", candidate_index))
    skill_score = _clamp(_feature_value(features, "skill_score", candidate_index))
    honeypot_penalty = _clamp(_feature_value(features, "honeypot_penalty", candidate_index))
    disqualifier_penalty = _clamp(_feature_value(features, "disqualifier_penalty", candidate_index))

    base_score = _clamp(
        JD_MATCH_WEIGHT * semantic_score
        + CAREER_WEIGHT * career_score
        + BEHAVIORAL_WEIGHT * behavioral_score
        + LOCATION_WEIGHT * location_score
        + SKILLS_WEIGHT * skill_score
    )
    final_score = _clamp(base_score * (1.0 - honeypot_penalty) * (1.0 - disqualifier_penalty))

    return {
        "semantic_similarity": semantic_score,
        "career_score": career_score,
        "behavioral_score": behavioral_score,
        "location_score": location_score,
        "skill_score": skill_score,
        "honeypot_penalty": honeypot_penalty,
        "disqualifier_penalty": disqualifier_penalty,
        "base_score": base_score,
        "final_score": final_score,
        "safe_score": _clean_score(guarded_score, default=FALLBACK_SCORE),
    }


def _feature_value(
    features: Mapping[str, npt.NDArray[np.float32]],
    feature_name: str,
    candidate_index: int,
) -> float:
    """Return one finite cached feature value for a candidate row."""
    try:
        value = features[feature_name][candidate_index]
    except IndexError as exc:
        raise ValueError(f"Feature {feature_name!r} missing candidate index {candidate_index}.") from exc
    return _finite_float(value, feature_name)


def _candidate_record_from_metadata(metadata: CandidateMetaCache, candidate_index: int) -> dict[str, Any]:
    """Build a lightweight fact-only candidate record for scoring and reasoning."""
    candidate_id = metadata["candidate_id"][candidate_index]
    current_title = metadata["current_title"][candidate_index]
    current_company = metadata["current_company"][candidate_index]
    years_of_experience = metadata["years_of_experience"][candidate_index]
    notice_period_days = metadata["notice_period_days"][candidate_index]
    recruiter_response_rate = metadata["recruiter_response_rate"][candidate_index]

    profile: dict[str, Any] = {}
    if current_title is not None:
        profile["current_title"] = current_title
    if current_company is not None:
        profile["current_company"] = current_company
    if years_of_experience is not None:
        profile["years_of_experience"] = years_of_experience

    redrob_signals: dict[str, Any] = {}
    if notice_period_days is not None:
        redrob_signals["notice_period_days"] = notice_period_days
    if recruiter_response_rate is not None:
        redrob_signals["recruiter_response_rate"] = recruiter_response_rate

    return {
        "candidate_id": candidate_id,
        "current_title": current_title,
        "current_company": current_company,
        "years_of_experience": years_of_experience,
        "profile": profile,
        "redrob_signals": redrob_signals,
    }


def _validate_scored_candidates(scored: Sequence[ScoredCandidate]) -> None:
    """Validate scored candidate IDs and raw scores before sorting."""
    if len(scored) < TOP_N:
        raise ValueError(f"At least {TOP_N} scored candidates are required, got {len(scored)}.")
    candidate_ids = [candidate.candidate_id for candidate in scored]
    scores = np.asarray([candidate.raw_score for candidate in scored], dtype=np.float32)
    normalizer.validate_ranking(candidate_ids, scores)


def _rank_top_candidates(
    scored_candidates: Sequence[ScoredCandidate],
) -> tuple[list[ScoredCandidate], npt.NDArray[np.float32]]:
    """Sort candidates by raw score and normalize the top-100 scores."""
    if len(scored_candidates) < TOP_N:
        raise ValueError(f"Need at least {TOP_N} scored candidates, got {len(scored_candidates)}.")

    sorted_candidates = sorted(scored_candidates, key=lambda candidate: candidate.raw_score, reverse=True)
    top_candidates = sorted_candidates[:TOP_N]
    raw_scores = np.asarray([candidate.raw_score for candidate in top_candidates], dtype=np.float32)
    candidate_ids = [candidate.candidate_id for candidate in top_candidates]

    normalizer.validate_ranking(candidate_ids, raw_scores)
    normalizer.validate_monotonic(raw_scores)
    normalized_scores = normalizer.normalize_scores(raw_scores)
    normalizer.validate_monotonic(normalized_scores)

    LOGGER.info(
        "Selected top candidates",
        extra={
            "top_n": TOP_N,
            "raw_min": float(np.min(raw_scores)),
            "raw_max": float(np.max(raw_scores)),
            "normalized_min": float(np.min(normalized_scores)),
            "normalized_max": float(np.max(normalized_scores)),
        },
    )
    return top_candidates, normalized_scores


def _build_submission_rows(
    top_candidates: Sequence[ScoredCandidate],
    normalized_scores: npt.NDArray[np.float32],
) -> list[SubmissionRow]:
    """Generate grounded reasoning and assemble final submission rows."""
    if len(top_candidates) != TOP_N:
        raise ValueError(f"Expected {TOP_N} top candidates, got {len(top_candidates)}.")
    if normalized_scores.shape != (TOP_N,):
        raise ValueError(f"Expected normalized score shape ({TOP_N},), got {normalized_scores.shape}.")

    rows: list[SubmissionRow] = []
    iterator = enumerate(zip(top_candidates, normalized_scores.tolist(), strict=True), start=1)
    for rank, (candidate, score) in tqdm_loader(iterator, total=TOP_N, desc="Generating reasoning"):
        clean_score = _clean_score(score, default=normalizer.DEFAULT_OUTPUT_MIN)
        reasoning = generate_reasoning(candidate.candidate_record, rank=rank, score=clean_score)
        rows.append(
            {
                "candidate_id": candidate.candidate_id,
                "rank": rank,
                "score": clean_score,
                "reasoning": reasoning,
            }
        )

    validate_reasoning_variance([row["reasoning"] for row in rows])
    return rows


def _validate_submission(rows: Sequence[SubmissionRow]) -> None:
    """Validate final submission invariants before writing CSV."""
    if len(rows) != TOP_N:
        raise ValueError(f"Submission must contain exactly {TOP_N} rows, got {len(rows)}.")

    candidate_ids: list[str] = []
    ranks: list[int] = []
    scores: list[float] = []
    for row_index, row in enumerate(rows):
        candidate_id = row.get("candidate_id")
        rank = row.get("rank")
        score = row.get("score")
        reasoning = row.get("reasoning")

        if not isinstance(candidate_id, str) or not candidate_id:
            raise ValueError(f"Submission row {row_index} has an invalid candidate_id.")
        if isinstance(rank, bool) or not isinstance(rank, Integral):
            raise ValueError(f"Submission row {row_index} has an invalid rank: {rank!r}.")
        clean_rank = int(rank)
        clean_score = _finite_float(score, f"submission score at row {row_index}")
        if not normalizer.DEFAULT_OUTPUT_MIN <= clean_score <= normalizer.DEFAULT_OUTPUT_MAX:
            raise ValueError(
                f"Submission score at row {row_index} is outside "
                f"[{normalizer.DEFAULT_OUTPUT_MIN}, {normalizer.DEFAULT_OUTPUT_MAX}]: {clean_score}."
            )
        if not isinstance(reasoning, str) or not reasoning.strip():
            raise ValueError(f"Submission row {row_index} has empty reasoning.")

        candidate_ids.append(candidate_id)
        ranks.append(clean_rank)
        scores.append(clean_score)

    expected_ranks = list(range(1, TOP_N + 1))
    if ranks != expected_ranks:
        raise ValueError(f"Submission ranks must be exactly 1-{TOP_N}.")
    if len(set(candidate_ids)) != len(candidate_ids):
        raise ValueError("Submission contains duplicate candidate IDs.")

    score_array = np.asarray(scores, dtype=np.float32)
    normalizer.validate_ranking(candidate_ids, score_array)
    normalizer.validate_monotonic(score_array)
    LOGGER.info("Submission validation complete", extra={"rows": len(rows)})


def _write_submission(rows: Sequence[SubmissionRow], path: pathlib.Path) -> None:
    """Write the validated submission CSV with UTF-8 encoding."""
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = path.with_name(path.name + ".tmp")
    with temp_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=CSV_COLUMNS, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    "candidate_id": row["candidate_id"],
                    "rank": row["rank"],
                    "score": f"{row['score']:.{SCORE_DECIMAL_PLACES}f}",
                    "reasoning": row["reasoning"],
                }
            )
    temp_path.replace(path)
    LOGGER.info("Wrote submission CSV", extra={"path": str(path), "rows": len(rows), "bytes": path.stat().st_size})


def _is_finite_real(value: object) -> bool:
    """Return whether a value is a non-bool finite real number."""
    return not isinstance(value, bool) and isinstance(value, Real) and np.isfinite(float(value))


def _finite_float(value: object, name: str) -> float:
    """Validate and return a finite float."""
    if not _is_finite_real(value):
        raise ValueError(f"{name} must be a finite real number, got {value!r}.")
    return float(value)


def _clean_score(value: object, default: float) -> float:
    """Return a finite clamped score or a finite clamped default."""
    try:
        return _clamp(_finite_float(value, "score"))
    except ValueError:
        return _clamp(float(default))


def _clamp(value: float) -> float:
    """Clamp a score or penalty to the inclusive 0-1 range."""
    return max(0.0, min(1.0, float(value)))


__all__ = ["main"]


if __name__ == "__main__":
    raise SystemExit(main())
